import uuid

import structlog
from flask import Blueprint, jsonify, request

from app.core.chunker import chunk_documents
from app.core.config import get_settings
from app.core.vector_store import get_vector_store
from app.middleware.auth import token_required
from app.middleware.rate_limit import rate_limit

logger = structlog.get_logger()
ingest_bp = Blueprint("ingest", __name__, url_prefix="/ingest")

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "html", "csv", "yaml", "yml", "json", "md"}


def extract_text(file, filename):
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    raw = file.read()
    if ext == "pdf":
        import io

        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext in ("docx",):
        import io

        from docx import Document
        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("latin-1")


@ingest_bp.route("", methods=["POST"])
@token_required
@rate_limit
def ingest_document():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "Empty filename"}), 400
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"Unsupported file type: .{ext}"}), 400
    try:
        text = extract_text(file, file.filename)
        doc_id = str(uuid.uuid4())[:8]
        from langchain_core.documents import Document
        doc_obj = Document(page_content=text, metadata={"filename": file.filename, "source": file.filename})
        chunks = chunk_documents([doc_obj], get_settings())
        store = get_vector_store()
        count = store.add_documents(chunks, doc_id=doc_id)
        return jsonify({"status": "indexed", "filename": file.filename, "doc_id": doc_id, "num_chunks": count}), 200
    except Exception as exc:
        logger.error("ingest_failed", error=str(exc), exc_info=True)
        return jsonify({"error": "Ingest failed: " + str(exc)}), 500


