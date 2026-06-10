# MIT License
# Copyright (c) 2026 RAG-QA Contributors

"""Document loaders for supported file formats."""

import csv
import io
from pathlib import Path

import structlog
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_community.document_loaders import CSVLoader, PyPDFLoader, TextLoader
from langchain_core.documents import Document
from pypdf import PdfReader

from app.core.config import Settings

logger = structlog.get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".csv"}


class UnsupportedFileTypeError(ValueError):
    """Raised when an uploaded file has an unsupported extension."""


def _load_pdf(file_path: Path) -> list[Document]:
    """Load a PDF file using PyPDFLoader.

    Args:
        file_path: Path to the PDF file.

    Returns:
        List of LangChain Document objects, one per page.
    """
    loader = PyPDFLoader(str(file_path))
    return loader.load()


def _load_txt(file_path: Path) -> list[Document]:
    """Load a plain-text file.

    Args:
        file_path: Path to the text file.

    Returns:
        List containing a single Document.
    """
    loader = TextLoader(str(file_path), encoding="utf-8")
    return loader.load()


def _load_html(file_path: Path) -> list[Document]:
    """Load an HTML file and extract visible text.

    Args:
        file_path: Path to the HTML file.

    Returns:
        List containing a single Document with extracted text.
    """
    raw_html = file_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw_html, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return [Document(page_content=text, metadata={"source": str(file_path)})]


def _load_docx(file_path: Path) -> list[Document]:
    """Load a DOCX file and concatenate paragraph text.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        List containing a single Document.
    """
    doc = DocxDocument(str(file_path))
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    text = "\n\n".join(paragraphs)
    return [Document(page_content=text, metadata={"source": str(file_path)})]


def _load_csv(file_path: Path) -> list[Document]:
    """Load a CSV file as row-level documents.

    Args:
        file_path: Path to the CSV file.

    Returns:
        List of Documents, one per CSV row.
    """
    loader = CSVLoader(str(file_path))
    return loader.load()


def _load_csv_from_bytes(content: bytes, filename: str) -> list[Document]:
    """Load CSV content from raw bytes.

    Args:
        content: Raw CSV bytes.
        filename: Original filename for metadata.

    Returns:
        List of Documents, one per CSV row.
    """
    text_content = content.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text_content))
    documents: list[Document] = []
    for row_index, row in enumerate(reader):
        row_text = ", ".join(f"{key}: {value}" for key, value in row.items())
        documents.append(
            Document(
                page_content=row_text,
                metadata={"source": filename, "row": row_index},
            )
        )
    return documents


def validate_extension(filename: str, settings: Settings) -> str:
    """Validate that the file extension is in the allowlist.

    Args:
        filename: Original filename.
        settings: Application settings with allowed extensions.

    Returns:
        Lowercase extension including the leading dot.

    Raises:
        UnsupportedFileTypeError: If the extension is not allowed.
    """
    extension = Path(filename).suffix.lower()
    if extension not in settings.allowed_extensions:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{extension}'. "
            f"Allowed: {', '.join(settings.allowed_extensions)}"
        )
    return extension


def load_document(file_path: Path, filename: str | None = None) -> list[Document]:
    """Load a document from disk based on its file extension.

    Args:
        file_path: Path to the file on disk.
        filename: Optional override for the source filename in metadata.

    Returns:
        List of LangChain Document objects.

    Raises:
        UnsupportedFileTypeError: If the file type is not supported.
        ValueError: If the file is empty or unreadable.
    """
    source_name = filename or file_path.name
    extension = file_path.suffix.lower()

    logger.info("loading_document", path=str(file_path), extension=extension)

    loaders = {
        ".pdf": _load_pdf,
        ".txt": _load_txt,
        ".html": _load_html,
        ".docx": _load_docx,
        ".csv": _load_csv,
    }

    loader_fn = loaders.get(extension)
    if loader_fn is None:
        raise UnsupportedFileTypeError(f"Unsupported file type: {extension}")

    documents = loader_fn(file_path)
    if not documents:
        raise ValueError(f"No content extracted from file: {source_name}")

    for doc in documents:
        doc.metadata["filename"] = source_name
        if "source" not in doc.metadata:
            doc.metadata["source"] = source_name

    logger.info("document_loaded", filename=source_name, pages=len(documents))
    return documents


def load_pdf_page_count(file_path: Path) -> int:
    """Return the number of pages in a PDF without full LangChain loading.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Number of pages, or 0 if not a PDF.
    """
    if file_path.suffix.lower() != ".pdf":
        return 0
    try:
        reader = PdfReader(str(file_path))
        return len(reader.pages)
    except Exception:
        return 0
